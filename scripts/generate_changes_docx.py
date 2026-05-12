#!/usr/bin/env python3
"""
国名・首都名の変更点をWordファイルに出力する
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import datetime

OUT_DIR = Path(__file__).resolve().parent.parent / "06_docs"
DOCX_PATH = OUT_DIR / "国名・首都名_変更一覧_20260512.docx"


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_note_paragraph(doc, label, text, label_color=None):
    """NOTE/IMPORTANT段落を追加"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run_label = p.add_run(f"【{label}】")
    run_label.bold = True
    run_label.font.size = Pt(9)
    if label_color:
        run_label.font.color.rgb = label_color

    run_text = p.add_run(f" {text}")
    run_text.font.size = Pt(9)


def build_document():
    doc = Document()

    # デフォルトフォント設定
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # ── タイトル ──
    title = doc.add_heading('国名・首都名 変更一覧', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)

    # サブタイトル
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('2025年5月28日付 検討会議資料（PDF）と外務省最新データの比較')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 比較情報
    info = doc.add_paragraph()
    info.paragraph_format.space_before = Pt(12)
    runs_data = [
        ("比較元：", True),
        ("国名・首都名検討会議資料（2025年5月28日付PDF）\n", False),
        ("比較先：", True),
        ("外務省ウェブサイト各国 data.html（2026年5月12日 Playwright取得）\n", False),
        ("比較日：", True),
        ("2026年5月12日", False),
    ]
    for text, bold in runs_data:
        r = info.add_run(text)
        r.bold = bold
        r.font.size = Pt(9)

    doc.add_paragraph()  # 空行

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 変更1: 赤道ギニア
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h = doc.add_heading('1. 赤道ギニア共和国 — 首都変更', level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    headers = ["項目", "2025年5月（PDF）", "2026年5月（外務省最新）"]
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # データ行
    data = [
        ["正式名称", "赤道ギニア共和国\n(Republic of Equatorial Guinea)", "赤道ギニア共和国\n(Republic of Equatorial Guinea)"],
        ["首都", "マラボ（Malabo）", "シウダ・デ・ラ・パス\n（Ciudad de la Paz）"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if r_idx == 1 and c_idx == 2:
                # 変更箇所をハイライト
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        r.bold = True

    add_note_paragraph(doc, "IMPORTANT",
        "2026年1月2日付の大統領令により、赤道ギニア政府は首都機能を従来のビオコ島マラボ（Malabo）から、"
        "大陸部ジブロオ州の新都市シウダ・デ・ラ・パス（Ciudad de la Paz）へ移転することを宣言した。"
        "ただし、実質的な首都機能は引き続きマラボに置かれている。",
        RGBColor(0xC0, 0x00, 0x00))

    add_note_paragraph(doc, "NOTE",
        "外務省の赤道ギニア基礎データ（令和8年3月6日更新）にも移転が明記されている。"
        "URL: https://www.mofa.go.jp/mofaj/area/eq_guinea/data.html",
        RGBColor(0x00, 0x70, 0xC0))

    doc.add_paragraph()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 変更2: ラトビア
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_heading('2. ラトビア共和国 — 首都名の表記変更', level=2)

    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h_text in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = h_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    data2 = [
        ["正式名称", "ラトビア共和国\n(Republic of Latvia)", "ラトビア共和国\n(Republic of Latvia)"],
        ["首都", "リガ", "リーガ"],
    ]
    for r_idx, row_data in enumerate(data2):
        for c_idx, cell_text in enumerate(row_data):
            cell = table2.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if r_idx == 1 and c_idx == 2:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        r.bold = True

    add_note_paragraph(doc, "IMPORTANT",
        "2026年（令和8年）4月1日より、ラトビア政府からの要請を踏まえ、"
        "ラトビア語の発音や言語・文化的背景に配慮する観点から、"
        "日本語表記が「リガ」から「リーガ」に変更された。",
        RGBColor(0xC0, 0x00, 0x00))

    add_note_paragraph(doc, "NOTE",
        "在ラトビア日本国大使館による告知: "
        "https://www.lv.emb-japan.go.jp/itpr_ja/11_000001_00547.html",
        RGBColor(0x00, 0x70, 0xC0))

    add_note_paragraph(doc, "NOTE",
        "これに伴い、在ラトビア日本国大使館および外務省において、"
        "公的な日本語表記が「リーガ」へと変更されている。",
        RGBColor(0x00, 0x70, 0xC0))

    doc.add_paragraph()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 変更3: リヒテンシュタイン
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_heading('3. リヒテンシュタイン公国 — 首都名の表記変更', level=2)

    table3 = doc.add_table(rows=3, cols=3)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h_text in enumerate(headers):
        cell = table3.rows[0].cells[i]
        cell.text = h_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    data3 = [
        ["正式名称", "リヒテンシュタイン公国\n(Principality of Liechtenstein)", "リヒテンシュタイン公国\n(Principality of Liechtenstein)"],
        ["首都", "ファドーツ", "ファドゥーツ"],
    ]
    for r_idx, row_data in enumerate(data3):
        for c_idx, cell_text in enumerate(row_data):
            cell = table3.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if r_idx == 1 and c_idx == 2:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        r.bold = True

    add_note_paragraph(doc, "NOTE",
        "ドイツ語原音「Vaduz」により近い日本語表記への変更。"
        "公式な変更通知や在外公館による告知は確認されていないが、"
        "外務省基礎データでの表記が「ファドゥーツ」に更新されている。"
        "なお、従来の「ファドーツ」表記は一部の書籍や地図で使用されていた。",
        RGBColor(0x00, 0x70, 0xC0))

    doc.add_paragraph()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 変更5: UAE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_heading('4. アラブ首長国連邦 — 英語名への略称追記', level=2)

    table5 = doc.add_table(rows=3, cols=3)
    table5.style = 'Table Grid'
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h_text in enumerate(headers):
        cell = table5.rows[0].cells[i]
        cell.text = h_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    data5 = [
        ["正式名称\n（英語）", "United Arab Emirates", "United Arab Emirates：UAE"],
        ["首都", "アブダビ", "アブダビ"],
    ]
    for r_idx, row_data in enumerate(data5):
        for c_idx, cell_text in enumerate(row_data):
            cell = table5.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if r_idx == 0 and c_idx == 2:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        r.bold = True

    add_note_paragraph(doc, "NOTE",
        "外務省 data.html 上の正式名称表記に略称「UAE」がコロン付きで追記された。"
        "国名そのものの変更ではなく、表記上の補足と思われる。"
        "公的な変更告知は確認されていない。",
        RGBColor(0x00, 0x70, 0xC0))

    # ── フッター ──
    doc.add_paragraph()
    doc.add_paragraph('―' * 40)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run(f'作成日: 2026年5月12日\n'
                        f'出典: 外務省「国・地域」data.html（Playwright headed モードで取得）')
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 保存
    doc.save(str(DOCX_PATH))
    print(f"✅ Word保存完了: {DOCX_PATH}")


if __name__ == "__main__":
    build_document()
